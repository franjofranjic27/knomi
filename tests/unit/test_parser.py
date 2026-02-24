"""Unit tests for knomi.ingest.parser."""

from pathlib import Path

from knomi.ingest.parser import _normalise, parse


def test_normalise_collapses_whitespace() -> None:
    assert _normalise("hello   world\t\nfoo") == "hello world foo"


def test_normalise_dehyphenates() -> None:
    assert _normalise("hyphen-\nated") == "hyphenated"


def test_normalise_strips_leading_trailing() -> None:
    assert _normalise("  hello  ") == "hello"


def test_parse_text_utf8(tmp_path: Path) -> None:
    f = tmp_path / "doc.txt"
    f.write_text("Hello world.", encoding="utf-8")
    result = parse(f)
    assert result == "Hello world."


def test_parse_text_md(tmp_path: Path) -> None:
    f = tmp_path / "doc.md"
    f.write_text("# Title\nSome content.", encoding="utf-8")
    result = parse(f)
    assert "Title" in result
    assert "Some content" in result


def test_parse_returns_empty_for_unknown_extension(tmp_path: Path) -> None:
    f = tmp_path / "data.xyz"
    f.write_text("irrelevant")
    assert parse(f) == ""


def test_parse_text_latin1_fallback(tmp_path: Path) -> None:
    f = tmp_path / "latin.txt"
    f.write_bytes("caf\xe9".encode("latin-1"))
    result = parse(f)
    assert "caf" in result


def test_parse_pdf(pdf_file: Path) -> None:
    result = parse(pdf_file)
    assert "Hello PDF world" in result
    assert len(result) > 0


def test_parse_docx(tmp_path: Path) -> None:
    from docx import Document  # type: ignore[import-untyped]

    path = tmp_path / "doc.docx"
    doc = Document()
    doc.add_paragraph("Hello DOCX world.")
    doc.add_paragraph("Second paragraph here.")
    doc.save(str(path))

    result = parse(path)
    assert "Hello DOCX world" in result
    assert "Second paragraph" in result


def test_parse_html(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        "<html><body><h1>Title</h1><p>HTML content here.</p></body></html>",
        encoding="utf-8",
    )
    result = parse(path)
    assert "Title" in result
    assert "HTML content here" in result
    assert "<h1>" not in result


def test_parse_html_strips_tags(tmp_path: Path) -> None:
    path = tmp_path / "styled.html"
    path.write_text(
        "<html><head><style>body{color:red}</style></head><body><p>Visible text.</p></body></html>",
        encoding="utf-8",
    )
    result = parse(path)
    assert "Visible text" in result
    assert "<style>" not in result
